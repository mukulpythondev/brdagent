import io
import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

class NumberedCanvas(canvas.Canvas):
    """
    Two-pass canvas to dynamically calculate the total page count
    and draw 'Page X of Y' in the footer along with headers.
    """
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        # We don't draw running header/footer on page 1 (the cover page)
        if self._pageNumber == 1:
            return
            
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#666666"))
        
        # Header
        self.drawString(54, 750, "BRD Forge — Generated Requirements Document")
        self.setStrokeColor(colors.HexColor("#E0E0E0"))
        self.setLineWidth(0.5)
        self.line(54, 742, 612 - 54, 742)
        
        # Footer
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(612 - 54, 36, page_str)
        self.drawString(54, 36, "Confidential — Useless AI BRD Forge")
        self.line(54, 48, 612 - 54, 48)
        
        self.restoreState()

def export_to_pdf(brd_data: dict) -> bytes:
    """
    Generate a professional PDF using ReportLab with custom cover page,
    running headers/footers, and clean tables for requirements.
    """
    pdf_buffer = io.BytesIO()
    
    # 0.5 inch margins = 36 pt, 0.75 inch = 54 pt. Page size is Letter (612 x 792 pt)
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    # Define custom styles
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=32,
        leading=38,
        textColor=colors.HexColor("#4285F4"), # Google Blue
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=16,
        leading=22,
        textColor=colors.HexColor("#666666"),
        spaceAfter=40
    )
    
    meta_style = ParagraphStyle(
        'CoverMeta',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=16,
        textColor=colors.HexColor("#333333"),
        spaceAfter=8
    )
    
    h1_style = ParagraphStyle(
        'BRDH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#4285F4"),
        spaceBefore=15,
        spaceAfter=10,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'BRDH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#2A2A3E"),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'BRDBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#333333"),
        spaceAfter=8
    )
    
    bullet_style = ParagraphStyle(
        'BRDBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )
    
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=11,
        textColor=colors.white
    )
    
    table_body_style = ParagraphStyle(
        'TableBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#333333")
    )
    
    table_body_bold_style = ParagraphStyle(
        'TableBodyBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#333333")
    )

    story = []
    
    # --- PAGE 1: COVER PAGE ---
    story.append(Spacer(1, 100))
    # Elegant logo line or border
    logo_data = [[""]]
    logo_table = Table(logo_data, colWidths=[504], rowHeights=[6])
    logo_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#4285F4")),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(logo_table)
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("BUSINESS REQUIREMENTS DOCUMENT", subtitle_style))
    story.append(Paragraph(brd_data.get("project_name", "Project Prototype").upper(), title_style))
    
    story.append(Spacer(1, 150))
    
    # Meta Section
    domain_name = brd_data.get("domain", "General")
    story.append(Paragraph(f"<b>Domain:</b> {domain_name}", meta_style))
    story.append(Paragraph(f"<b>Created:</b> {datetime.datetime.now().strftime('%B %d, %Y')}", meta_style))
    story.append(Paragraph("<b>Generated By:</b> BRD Forge (Intelligent Multi-Modal Agent)", meta_style))
    story.append(Paragraph("<b>Author Team:</b> Useless AI", meta_style))
    
    story.append(PageBreak())
    
    # --- PAGE 2+: CONTENT ---
    
    # 1. Executive Summary
    story.append(Paragraph("1. Executive Summary", h1_style))
    summary_text = brd_data.get("executive_summary", "No summary provided.")
    story.append(Paragraph(summary_text, body_style))
    story.append(Spacer(1, 12))
    
    # 2. Objectives
    story.append(Paragraph("2. Project Objectives", h1_style))
    objectives = brd_data.get("objectives", [])
    if not objectives:
        story.append(Paragraph("No objectives specified.", body_style))
    for i, obj in enumerate(objectives):
        story.append(Paragraph(f"<b>{i+1}.</b> {obj}", bullet_style))
    story.append(Spacer(1, 12))
    
    # 3. Scope
    story.append(Paragraph("3. Scope Description", h1_style))
    scope = brd_data.get("scope", {})
    in_scope = scope.get("in_scope", [])
    out_scope = scope.get("out_scope", [])
    
    story.append(Paragraph("<b>In-Scope Items:</b>", h2_style))
    if not in_scope:
        story.append(Paragraph("None specified.", body_style))
    for item in in_scope:
        story.append(Paragraph(f"• {item}", bullet_style))
        
    story.append(Paragraph("<b>Out-of-Scope Items:</b>", h2_style))
    if not out_scope:
        story.append(Paragraph("None specified.", body_style))
    for item in out_scope:
        story.append(Paragraph(f"• {item}", bullet_style))
    story.append(Spacer(1, 12))
    
    # 4. Stakeholders
    story.append(Paragraph("4. Key Stakeholders", h1_style))
    stakeholders = brd_data.get("stakeholders", [])
    if not stakeholders:
        story.append(Paragraph("No stakeholders identified.", body_style))
    for sh in stakeholders:
        story.append(Paragraph(f"• {sh}", bullet_style))
    story.append(Spacer(1, 12))
    
    # Page Break before requirements
    story.append(PageBreak())
    
    # 5. Functional Requirements
    story.append(Paragraph("5. Functional Requirements", h1_style))
    fr_list = brd_data.get("functional_requirements", [])
    if not fr_list:
        story.append(Paragraph("No functional requirements specified.", body_style))
    else:
        # Table columns: ID (60 pt) | Title (100 pt) | Description (180 pt) | Source (60 pt) | Conf/Priority (104 pt)
        table_data = [[
            Paragraph("ID", table_header_style),
            Paragraph("Title", table_header_style),
            Paragraph("Description", table_header_style),
            Paragraph("Source", table_header_style),
            Paragraph("Meta", table_header_style)
        ]]
        
        for fr in fr_list:
            fid = fr.get("id", "FR")
            title = fr.get("title", "")
            desc = fr.get("description", "")
            src = fr.get("source", "")
            conf = fr.get("confidence", "HIGH")
            priority = fr.get("priority", "MUST HAVE")
            conflict = fr.get("conflict")
            
            meta_txt = f"Conf: <b>{conf}</b><br/>Pri: {priority}"
            if conflict:
                meta_txt += "<br/><font color='red'><b>⚠️ CONFLICT</b></font>"
                
            table_data.append([
                Paragraph(fid, table_body_bold_style),
                Paragraph(title, table_body_bold_style),
                Paragraph(desc, table_body_style),
                Paragraph(src, table_body_style),
                Paragraph(meta_txt, table_body_style)
            ])
            
        t = Table(table_data, colWidths=[45, 95, 194, 80, 90])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4285F4")),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F7FA")]),
        ]))
        story.append(t)
        
    story.append(Spacer(1, 15))
    
    # 6. Non-Functional Requirements
    story.append(Paragraph("6. Non-Functional Requirements", h1_style))
    nfr_list = brd_data.get("non_functional_requirements", [])
    if not nfr_list:
        story.append(Paragraph("No non-functional requirements specified.", body_style))
    else:
        table_data = [[
            Paragraph("ID", table_header_style),
            Paragraph("Title", table_header_style),
            Paragraph("Description", table_header_style),
            Paragraph("Source", table_header_style),
            Paragraph("Confidence", table_header_style)
        ]]
        
        for nfr in nfr_list:
            nid = nfr.get("id", "NFR")
            title = nfr.get("title", "")
            desc = nfr.get("description", "")
            src = nfr.get("source", "")
            conf = nfr.get("confidence", "HIGH")
            
            table_data.append([
                Paragraph(nid, table_body_bold_style),
                Paragraph(title, table_body_bold_style),
                Paragraph(desc, table_body_style),
                Paragraph(src, table_body_style),
                Paragraph(conf, table_body_style)
            ])
            
        t = Table(table_data, colWidths=[45, 95, 224, 70, 70])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4285F4")),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F5F7FA")]),
        ]))
        story.append(t)
        
    story.append(Spacer(1, 15))
    
    # 7. Assumptions & Dependencies
    story.append(Paragraph("7. Assumptions & Dependencies", h1_style))
    assumptions = brd_data.get("assumptions", [])
    if not assumptions:
        story.append(Paragraph("No assumptions specified.", body_style))
    for asn in assumptions:
        story.append(Paragraph(f"• {asn}", bullet_style))
    story.append(Spacer(1, 12))
    
    # 8. Risks & Mitigations
    story.append(Paragraph("8. Risks & Mitigations", h1_style))
    risks = brd_data.get("risks", [])
    if not risks:
        story.append(Paragraph("No risks specified.", body_style))
    else:
        table_data = [[
            Paragraph("Risk Element", table_header_style),
            Paragraph("Suggested Mitigation Strategy", table_header_style)
        ]]
        for item in risks:
            risk = item.get("risk", "")
            mit = item.get("mitigation", "")
            table_data.append([
                Paragraph(risk, table_body_bold_style),
                Paragraph(mit, table_body_style)
            ])
        t = Table(table_data, colWidths=[240, 264])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#4285F4")),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
            ('TOPPADDING', (0,0), (-1,-1), 6),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t)
    story.append(Spacer(1, 12))
    
    # 9. Acceptance Criteria
    story.append(Paragraph("9. Acceptance Criteria", h1_style))
    ac = brd_data.get("acceptance_criteria", [])
    if not ac:
        story.append(Paragraph("No explicit acceptance criteria specified.", body_style))
    for criteria in ac:
        story.append(Paragraph(f"⬜  {criteria}", bullet_style))
        
    doc.build(story, canvasmaker=NumberedCanvas)
    pdf_bytes = pdf_buffer.getvalue()
    pdf_buffer.close()
    return pdf_bytes

def set_cell_background(cell, fill_hex):
    """
    Helper to color Word table cells.
    """
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def export_to_word(brd_data: dict) -> bytes:
    """
    Generate a professional Word Document using python-docx.
    """
    doc = Document()
    
    # Styles
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Section
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("BUSINESS REQUIREMENTS DOCUMENT")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(102, 102, 102)
    title_p.paragraph_format.space_after = Pt(2)
    
    proj_p = doc.add_paragraph()
    proj_run = proj_p.add_run(brd_data.get("project_name", "Project Prototype").upper())
    proj_run.bold = True
    proj_run.font.name = "Arial"
    proj_run.font.size = Pt(28)
    proj_run.font.color.rgb = RGBColor(66, 133, 244) # Google Blue
    proj_p.paragraph_format.space_after = Pt(24)
    
    # Metadata Block
    doc.add_paragraph(f"Domain: {brd_data.get('domain', 'General')}")
    doc.add_paragraph(f"Created: {datetime.datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("Author Team: Useless AI")
    doc.add_paragraph("Generated By: BRD Forge")
    
    doc.add_page_break()
    
    # Helper to add H1
    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(66, 133, 244)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p
        
    # Helper to add H2
    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 51, 51)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return p
        
    # Helper to add body
    def add_body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        return p
        
    # Executive Summary
    add_h1("1. Executive Summary")
    add_body(brd_data.get("executive_summary", ""))
    
    # Objectives
    add_h1("2. Project Objectives")
    for i, obj in enumerate(brd_data.get("objectives", [])):
        p = doc.add_paragraph(style='List Number')
        p.add_run(obj)
        p.paragraph_format.space_after = Pt(4)
        
    # Scope
    add_h1("3. Scope Description")
    add_h2("In-Scope")
    for item in brd_data.get("scope", {}).get("in_scope", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)
        
    add_h2("Out-of-Scope")
    for item in brd_data.get("scope", {}).get("out_scope", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)
        
    # Stakeholders
    add_h1("4. Key Stakeholders")
    for sh in brd_data.get("stakeholders", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(sh)
        p.paragraph_format.space_after = Pt(4)
        
    # Page break for requirements tables
    doc.add_page_break()
    
    # Functional Requirements
    add_h1("5. Functional Requirements")
    fr_list = brd_data.get("functional_requirements", [])
    if not fr_list:
        add_body("No functional requirements specified.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Description'
        hdr_cells[3].text = 'Source'
        hdr_cells[4].text = 'Meta Details'
        
        # Style headers
        for cell in hdr_cells:
            set_cell_background(cell, "4285F4")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
        for fr in fr_list:
            row_cells = table.add_row().cells
            row_cells[0].text = fr.get("id", "")
            row_cells[1].text = fr.get("title", "")
            row_cells[2].text = fr.get("description", "")
            row_cells[3].text = fr.get("source", "")
            
            meta_str = f"Confidence: {fr.get('confidence', '')}\nPriority: {fr.get('priority', '')}"
            if fr.get("conflict"):
                meta_str += f"\n[CONFLICT]: {fr.get('conflict')}"
            row_cells[4].text = meta_str
            
            # Shading alternate rows or styling conflicts
            if fr.get("conflict"):
                set_cell_background(row_cells[4], "FCE8E6") # Light red
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Non-Functional Requirements
    add_h1("6. Non-Functional Requirements")
    nfr_list = brd_data.get("non_functional_requirements", [])
    if not nfr_list:
        add_body("No non-functional requirements specified.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Description'
        hdr_cells[3].text = 'Source'
        hdr_cells[4].text = 'Confidence'
        
        # Style headers
        for cell in hdr_cells:
            set_cell_background(cell, "4285F4")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
        for nfr in nfr_list:
            row_cells = table.add_row().cells
            row_cells[0].text = nfr.get("id", "")
            row_cells[1].text = nfr.get("title", "")
            row_cells[2].text = nfr.get("description", "")
            row_cells[3].text = nfr.get("source", "")
            row_cells[4].text = nfr.get("confidence", "")
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Assumptions
    add_h1("7. Assumptions & Dependencies")
    for asn in brd_data.get("assumptions", []):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(asn)
        p.paragraph_format.space_after = Pt(4)
        
    # Risks
    add_h1("8. Risks & Mitigations")
    risks = brd_data.get("risks", [])
    if not risks:
        add_body("No risks specified.")
    else:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Risk Element'
        hdr_cells[1].text = 'Suggested Mitigation Strategy'
        for cell in hdr_cells:
            set_cell_background(cell, "4285F4")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
        for item in risks:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get("risk", "")
            row_cells[1].text = item.get("mitigation", "")
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Acceptance Criteria
    add_h1("9. Acceptance Criteria")
    for criteria in brd_data.get("acceptance_criteria", []):
        p = doc.add_paragraph()
        p.add_run("[  ]  " + criteria)
        p.paragraph_format.space_after = Pt(4)
        
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_bytes = word_buffer.getvalue()
    word_buffer.close()
    return word_bytes
